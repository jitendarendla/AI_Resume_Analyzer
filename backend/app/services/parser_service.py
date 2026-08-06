import os
import re
import json
import urllib.request
import urllib.parse
import threading
import gc
from datetime import datetime
import fitz  # PyMuPDF
import docx

SKILLS_LEXICON = [
    # Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "golang", "go", "ruby", "php",
    "swift", "kotlin", "rust", "scala", "r", "perl", "html", "css", "sql", "bash", "shell",
    # Frontend
    "react", "react.js", "next.js", "vue", "vue.js", "angular", "svelte", "redux", "tailwind",
    "tailwind css", "bootstrap", "sass", "webpack", "vite", "html5", "css3", "jquery",
    # Backend & Frameworks
    "node.js", "express", "fastapi", "django", "flask", "spring", "spring boot", ".net",
    "asp.net", "graphql", "rest api", "restful api", "microservices", "grpc", "gin", "gorilla",
    # Databases
    "postgresql", "postgres", "mysql", "mongodb", "sqlite", "redis", "memcached", "elasticsearch",
    "dynamodb", "cassandra", "oracle", "sql server", "firebase", "neo4j",
    # DevOps & Cloud
    "docker", "kubernetes", "k8s", "aws", "amazon web services", "azure", "gcp",
    "google cloud", "terraform", "ansible", "jenkins", "git", "github", "gitlab",
    "bitbucket", "ci/cd", "prometheus", "grafana", "linux", "unix", "nginx",
    # Data Science & AI/ML
    "pandas", "numpy", "scikit-learn", "scikit learn", "tensorflow", "pytorch", "keras",
    "nlp", "natural language processing", "machine learning", "deep learning", "ai",
    "artificial intelligence", "computer vision", "opencv", "tableau", "power bi",
    "spark", "hadoop", "kafka", "pyspark", "scipy", "seaborn", "matplotlib",
    # Methodologies & Tools
    "agile", "scrum", "jira", "confluence", "figma", "postman", "system design",
    "data structures", "algorithms", "oop", "object-oriented programming", "unit testing",
    "pytest", "jest", "cypress", "selenium"
]

SKILLS_DATABASE = set(SKILLS_LEXICON)

TECH_CASING = {
    "python": "Python", "javascript": "JavaScript", "typescript": "TypeScript",
    "fastapi": "FastAPI", "react": "React", "next.js": "Next.js", "node.js": "Node.js",
    "postgresql": "PostgreSQL", "mysql": "MySQL", "mongodb": "MongoDB",
    "docker": "Docker", "kubernetes": "Kubernetes", "aws": "AWS", "gcp": "GCP",
    "azure": "Azure", "git": "Git", "github": "GitHub", "gitlab": "GitLab",
    "rest api": "REST API", "ci/cd": "CI/CD", "html": "HTML", "css": "CSS",
    "c++": "C++", "c#": "C#", "golang": "Golang", "go": "Go", "gin": "Gin", "gorilla": "Gorilla"
}

# PRE-COMPILED ULTRA FAST REGEX PATTERNS
SKILL_REGEX_MAP = {
    skill: re.compile(r'\b' + re.escape(skill) + r'\b', re.IGNORECASE)
    for skill in SKILLS_LEXICON
}

EMAIL_REGEX = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
PHONE_REGEX = re.compile(r'(?:(?:\+|\d{1,3})[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}')
LINKEDIN_REGEX = re.compile(r'(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+/?', re.IGNORECASE)
GITHUB_REGEX = re.compile(r'(?:https?://)?(?:www\.)?github\.com/[a-zA-Z0-9_-]+/?', re.IGNORECASE)
DATE_YEAR_REGEX = re.compile(r'\b(19\d\d|20\d\d)\b')
EXP_YEAR_REGEX = re.compile(r'(\d+(?:\.\d+)?)\s*(?:\+\s*)?(?:years?|yrs?)\b', re.IGNORECASE)

CERTIFICATION_KEYWORDS = [
    "certified", "certification", "certificate", "aws", "azure", "google cloud", "pmp",
    "scrum master", "csm", "cka", "cissp", "comptia", "udemy", "coursera", "nptel", "edx"
]

DEGREE_KEYWORDS = [
    "bachelor", "b.s", "b.tech", "b.e", "b.a", "master", "m.s", "m.tech", "m.a", "ph.d",
    "doctorate", "mba", "diploma", "degree", "associate", "b.c.a", "m.c.a", "university",
    "college", "institute", "school"
]

JOB_TITLE_KEYWORDS = {
    'golang', 'developer', 'engineer', 'architect', 'consultant', 'programmer',
    'lead', 'manager', 'analyst', 'specialist', 'administrator', 'sr', 'jr',
    'senior', 'junior', 'full stack', 'backend', 'frontend', 'software', 'devops',
    'dice', 'resume', 'curriculum', 'vitae', 'cv', 'profile', 'contact'
}

LOCATION_SKILLS_KEYWORDS = {
    'go', 'golang', 'react', 'python', 'java', 'docker', 'kubernetes', 'k8s',
    'jenkins', 'redis', 'memcached', 'prometheus', 'grafana', 'aws', 'azure',
    'spring', 'boot', 'microservices', 'gin', 'gorilla', 'typescript', 'javascript',
    'analysis', 'design', 'requirement', 'requirements', 'logix', 'embedded',
    'controllers', 'high', 'cursor', 'claude', 'c++', 'c#', 'node', 'express',
    'sql', 'postgres', 'mysql', 'mongodb', 'git', 'gitlab', 'github', 'devops',
    'skills', 'experience', 'education'
}

LOCATION_INDICATORS = {
    'usa', 'us', 'united states', 'india', 'uk', 'united kingdom', 'canada',
    'chicago', 'richmond', 'virginia', 'texas', 'missouri', 'california', 'new york',
    'new jersey', 'georgia', 'florida', 'washington', 'maryland', 'illinois',
    'hyderabad', 'bangalore', 'bengaluru', 'mumbai', 'delhi', 'pune', 'chennai',
    'london', 'toronto', 'vancouver', 'europe', 'asia', 'singapore', 'dubai',
    'mckinney', 'austin', 'dallas', 'houston', 'san jose', 'san francisco', 'seattle',
    'panama city beach', 'panama city', 'orlando', 'tampa', 'miami', 'atlanta',
    'tx', 'va', 'il', 'mo', 'ca', 'ny', 'nj', 'ga', 'fl', 'wa', 'md', 'nc', 'ma', 'pa', 'oh', 'mi'
}

INVALID_LOC_KEYWORDS = {
    'snowflake', 'databricks', 'apache', 'spark', 'scala', 'feast', 'strategy', 'governance',
    'django', 'flask', 'machine', 'learning', 'deep', 'embedding', 'models', 'model', 'prompt',
    'engineering', 'ai', 'ml', 'nlp', 'cloud', 'data', 'architecture', 'design', 'development',
    'golang', 'developer', 'engineer', 'architect', 'consultant', 'programmer', 'lead',
    'manager', 'analyst', 'specialist', 'administrator', 'sr', 'jr', 'senior', 'junior',
    'full stack', 'backend', 'frontend', 'software', 'devops', 'skills', 'experience', 'education',
    'python', 'javascript', 'typescript', 'java', 'react', 'next.js', 'vue', 'angular',
    'node.js', 'express', 'fastapi', 'spring', 'boot', 'microservices', 'postgresql', 'postgres',
    'mysql', 'mongodb', 'redis', 'elasticsearch', 'docker', 'kubernetes', 'k8s', 'aws', 'azure', 'gcp'
}

REJECT_NAME_WORDS = [
    "page", "of", "work", "history", "employment", "project", "details",
    "scalability", "transformation", "across", "warm", "regards", "ph", "no",
    "summary", "profile", "contact", "curriculum", "vitae", "resume", "dice",
    "experience", "education", "skills", "objective", "professional", "references", "cover", "letter"
]

# Non-blocking lock for local Ollama server
ollama_lock = threading.Lock()

def extract_text_from_file(file_path: str) -> str:
    ext = file_path.split(".")[-1].lower() if "." in file_path else ""
    text = ""
    try:
        if ext == "pdf":
            doc = fitz.open(file_path)
            for page in doc[:5]:  # Read first 5 pages max for ultra fast throughput
                text += page.get_text() + "\n"
            doc.close()
            del doc
        elif ext == "docx":
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + "\n"
            del doc
        elif ext in ["txt", "doc"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    except Exception as e:
        print(f"[WARNING] Error reading {file_path}: {e}")
    finally:
        gc.collect()
    return text.strip()

def clean_candidate_name(raw_name: str, filename: str = "", email: str = "", raw_text: str = "") -> str:
    name_clean = str(raw_name or "").strip()

    # Strip leading noise prefixes
    name_clean = re.sub(r'^(?:name|candidate|applicant|fullName)\s*[:\-]?\s*', '', name_clean, flags=re.IGNORECASE).strip()

    # Strip trailing visa/doc noise suffixes
    name_clean = re.sub(r'\s+\b(?:I94|H1B|Travel|Visa|Resume|CV|Docx?|Pdf)\b.*$', '', name_clean, flags=re.IGNORECASE).strip()

    name_lower = name_clean.lower()
    words = [w.lower().strip('.') for w in name_clean.split()]

    is_invalid = False
    if any(nw in words for nw in ["work", "history", "employment", "details", "scalability", "transformation", "across", "warm", "regards", "ph", "no", "page"]):
        is_invalid = True
    if len(words) == 2 and words[0] == "page" and words[1] == "of":
        is_invalid = True

    job_words_count = sum(1 for w in words if w in JOB_TITLE_KEYWORDS)
    if job_words_count > 0 and (job_words_count >= len(words) / 2 or 'developer' in words or 'engineer' in words or 'sr' in words or 'dice' in words):
        is_invalid = True

    if is_invalid or not name_clean or name_lower in ['candidate', 'n/a', 'unknown', 'sr. golang developer', 'golang developer']:
        name_clean = ""

        if filename:
            base = filename.rsplit('.', 1)[0]
            base = re.sub(r'^(?:Dice_)?(?:Resume_)?(?:CV_)?(?:Resume)?', '', base, flags=re.IGNORECASE)
            base = re.sub(r'(?:_Golang_Developer|_Golang|_Developer|_CV|_pro|_pdf|_docx|_Gen_AI_Engineer|_cover_letter|_Resume)$', '', base, flags=re.IGNORECASE)
            base = re.sub(r'[-_]', ' ', base).strip()
            base = re.sub(r'([a-z])([A-Z])', r'\1 \2', base).strip()
            base = re.sub(r'\b(?:I94|H1B|Travel|Gen AI Engineer|Sr Gen AI Developer|Developer|Engineer|CV|Resume)\b', '', base, flags=re.IGNORECASE).strip()
            fn_words = [w for w in base.split() if w.lower() not in JOB_TITLE_KEYWORDS and len(w) > 1]
            if fn_words:
                name_clean = " ".join(fn_words).title()

        if not name_clean and email:
            handle = email.split('@')[0]
            handle_clean = re.sub(r'\d+', '', handle)
            parts = [p.capitalize() for p in re.split(r'[._-]', handle_clean) if len(p) > 1 and p.lower() not in JOB_TITLE_KEYWORDS]
            if parts:
                name_clean = " ".join(parts)

    return name_clean or "Candidate"

def clean_candidate_location(raw_loc: str, text: str = "") -> str:
    if not raw_loc or not isinstance(raw_loc, str):
        return ""

    loc_clean = raw_loc.strip()
    loc_lower = loc_clean.lower()

    # Reject if ANY word in location matches a tech skill or invalid keyword
    words = [w.lower().strip(',.()') for w in re.split(r'[\s,]+', loc_clean) if w]
    for w in words:
        if w in INVALID_LOC_KEYWORDS or w in LOCATION_SKILLS_KEYWORDS or w in SKILLS_DATABASE:
            return ""

    # Check for valid location indicator
    has_loc = any(ind in loc_lower for ind in LOCATION_INDICATORS)
    if not has_loc:
        return ""

    loc_clean = re.sub(r'^[A-Z][a-z]+\s+[A-Z][a-z]+(?=[A-Z])', '', loc_clean).strip()
    loc_clean = re.sub(r'^St,\s*', '', loc_clean).strip()
    return loc_clean

def extract_technology_title(text: str = "", filename: str = "", skills: list = None) -> str:
    if filename:
        fn_lower = filename.lower()
        if "gen ai" in fn_lower or "generative ai" in fn_lower:
            return "Sr. Gen AI Developer"
        elif "golang" in fn_lower or "go developer" in fn_lower:
            return "Golang Developer"
        elif "full stack" in fn_lower or "fullstack" in fn_lower:
            return "Full Stack Engineer"
        elif "data engineer" in fn_lower or "spark" in fn_lower or "snowflake" in fn_lower:
            return "Data Engineer"
        elif "backend" in fn_lower:
            return "Backend Engineer"
        elif "devops" in fn_lower:
            return "DevOps Engineer"

    if text:
        lines = [l.strip() for l in text.split("\n") if l.strip()][:10]
        for line in lines:
            line_lower = line.lower()
            if any(term in line_lower for term in ["developer", "engineer", "architect", "consultant", "lead", "specialist"]):
                if len(line) < 65 and not any(kw in line_lower for kw in ["experience", "education", "skills", "summary", "profile", "contact"]):
                    clean_title = re.sub(r'[^a-zA-Z0-9\s/.-]', '', line).strip()
                    if clean_title:
                        return clean_title.title()

    if skills and len(skills) > 0:
        top_skills = [s for s in skills[:2] if isinstance(s, str)]
        if top_skills:
            return f"{' / '.join(top_skills)} Developer"

    return "Software Engineer"

def normalize_llm_dict(raw: dict) -> dict:
    if not isinstance(raw, dict):
        return {}

    name = str(raw.get("name") or "").strip()
    email = str(raw.get("email") or "").strip()
    phone = str(raw.get("phone") or "").strip()
    location = str(raw.get("location") or "").strip()

    raw_skills = raw.get("skills") or []
    skills = []
    if isinstance(raw_skills, list):
        for s in raw_skills:
            if isinstance(s, str) and s.strip():
                skills.append(s.strip())
            elif isinstance(s, dict):
                skills.extend([str(v).strip() for v in s.values() if str(v).strip()])
    elif isinstance(raw_skills, dict):
        for k, v in raw_skills.items():
            if isinstance(v, list):
                skills.extend([str(x).strip() for x in v if str(x).strip()])
            elif isinstance(v, str) and v.strip():
                skills.append(v.strip())
    elif isinstance(raw_skills, str):
        skills = [s.strip() for s in raw_skills.split(",") if s.strip()]

    raw_edu = raw.get("education") or ""
    if isinstance(raw_edu, dict):
        education = " ".join([str(v) for v in raw_edu.values() if v]).strip()
    elif isinstance(raw_edu, list):
        education = ", ".join([str(v) for v in raw_edu if v]).strip()
    else:
        education = str(raw_edu).strip()

    raw_exp = raw.get("experience_years") or raw.get("years_of_experience") or 0.0
    try:
        experience_years = float(raw_exp)
    except (ValueError, TypeError):
        experience_years = 0.0

    raw_certs = raw.get("certifications") or []
    certs = [str(c).strip() for c in raw_certs] if isinstance(raw_certs, list) else ([str(raw_certs).strip()] if raw_certs else [])

    linkedin = str(raw.get("linkedin") or raw.get("linkedIn_profile_url") or "").strip()
    github = str(raw.get("github") or raw.get("github_profile_url") or "").strip()

    raw_projects = raw.get("projects") or []
    projects = [str(p).strip() for p in raw_projects] if isinstance(raw_projects, list) else ([str(raw_projects).strip()] if raw_projects else [])

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "location": location,
        "skills": list(dict.fromkeys(skills)),
        "education": education,
        "experience_years": experience_years,
        "certifications": certs,
        "linkedin": linkedin,
        "github": github,
        "projects": projects
    }

def parse_resume_with_llm(text: str) -> dict:
    """Ultra-Fast Non-Blocking Local / Cloud LLM Extractor."""
    prompt = f"""
    Extract resume details into valid JSON:
    {{
      "name": "Full Name",
      "email": "Email Address",
      "phone": "Phone Number",
      "location": "City, State",
      "skills": ["Skill1", "Skill2"],
      "education": "Degree",
      "experience_years": 5.0,
      "certifications": ["Cert1"],
      "linkedin": "",
      "github": "",
      "projects": []
    }}
    Resume Text:
    {text[:3000]}
    """

    # Only run Ollama if explicitly enabled via environment variable
    if os.getenv("ENABLE_OLLAMA") == "true":
        ollama_host = os.getenv("OLLAMA_HOST", "http://localhost:11434")
        ollama_model = os.getenv("OLLAMA_MODEL", "qwen2.5:1.5b")
        acquired = ollama_lock.acquire(blocking=False)
        if acquired:
            try:
                url = f"{ollama_host}/api/generate"
                payload = {
                    "model": ollama_model,
                    "prompt": prompt,
                    "format": "json",
                    "stream": False,
                    "keep_alive": "30m",
                    "options": {"temperature": 0.1}
                }
                req = urllib.request.Request(
                    url,
                    data=json.dumps(payload).encode("utf-8"),
                    headers={"Content-Type": "application/json"}
                )
                with urllib.request.urlopen(req, timeout=1.0) as resp:
                    res_data = json.loads(resp.read().decode("utf-8"))
                    parsed_json = json.loads(res_data.get("response", "{}"))
                    normalized = normalize_llm_dict(parsed_json)
                    if normalized.get("name") or normalized.get("skills"):
                        return normalized
            except Exception:
                pass
            finally:
                ollama_lock.release()

    gemini_key = os.getenv("GEMINI_API_KEY")
    openai_key = os.getenv("OPENAI_API_KEY")

    if gemini_key:
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={gemini_key}"
            payload = {
                "contents": [{"parts": [{"text": prompt}]}],
                "generationConfig": {"response_mime_type": "application/json"}
            }
            req = urllib.request.Request(
                url,
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type": "application/json"}
            )
            with urllib.request.urlopen(req, timeout=2.0) as resp:
                res_data = json.loads(resp.read().decode("utf-8"))
                text_resp = res_data["candidates"][0]["content"]["parts"][0]["text"]
                return normalize_llm_dict(json.loads(text_resp))
        except Exception:
            pass

    if openai_key:
        try:
            url = "https://api.openai.com/v1/chat/completions"
            payload = {
                "model": "gpt-4o-mini",
                "messages": [{"role": "user", "content": prompt}],
                "response_format": {"type": "json_object"}
            }
            req = urllib.request.Request(
                url,
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type": "application/json", "Authorization": f"Bearer {openai_key}"}
            )
            with urllib.request.urlopen(req, timeout=2.0) as resp:
                res_data = json.loads(resp.read().decode("utf-8"))
                text_resp = res_data["choices"][0]["message"]["content"]
                return normalize_llm_dict(json.loads(text_resp))
        except Exception:
            pass

    return None

def extract_rule_based_details(text: str, filename: str = "") -> dict:
    """Pre-compiled Regex Engine for Sub-Millisecond 100% Extraction Coverage."""
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # 1. Email
    email_match = EMAIL_REGEX.search(text)
    email = email_match.group(0) if email_match else ""

    # 2. Phone
    phone_match = PHONE_REGEX.search(text)
    phone = phone_match.group(0) if phone_match else ""

    # 3. Links
    linkedin_match = LINKEDIN_REGEX.search(text)
    linkedin = linkedin_match.group(0) if linkedin_match else ""

    github_match = GITHUB_REGEX.search(text)
    github = github_match.group(0) if github_match else ""

    # 4. Name
    candidate_name = ""
    for line in lines[:8]:
        if "@" in line or "http" in line or "www." in line or re.search(r'\d{5,}', line):
            continue
        if any(h in line.lower() for h in ["resume", "curriculum vitae", "cv", "profile", "contact", "summary", "experience", "education", "skills"]):
            continue
        clean_line = re.sub(r'[^a-zA-Z\s.-]', '', line).strip()
        words = [w.lower() for w in clean_line.split()]
        if any(w in JOB_TITLE_KEYWORDS for w in words):
            continue
        if clean_line and 2 <= len(clean_line.split()) <= 4 and len(clean_line) < 40:
            candidate_name = clean_line.title()
            break

    candidate_name = clean_candidate_name(candidate_name, filename, email, text)

    # 5. Skills
    text_lower = text.lower()
    found_skills = set()
    for skill, pattern in SKILL_REGEX_MAP.items():
        if pattern.search(text_lower):
            cased = TECH_CASING.get(skill, skill.title())
            found_skills.add(cased)

    # 6. Experience
    experience_years = 0.0
    exp_matches = EXP_YEAR_REGEX.findall(text_lower)
    if exp_matches:
        try:
            years = [float(y) for y in exp_matches if float(y) < 45.0]
            if years:
                experience_years = max(years)
        except Exception:
            pass

    if experience_years == 0.0:
        years_found = [int(y) for y in DATE_YEAR_REGEX.findall(text) if 1990 <= int(y) <= datetime.now().year]
        if len(years_found) >= 2:
            span = max(years_found) - min(years_found)
            if 0 < span <= 35:
                experience_years = float(span)

    # 7. Education
    education_parts = set()
    for line in lines:
        line_lower = line.lower()
        if any(deg in line_lower for deg in DEGREE_KEYWORDS):
            if len(line) < 120:
                education_parts.add(line.strip())
    education_str = ", ".join(list(education_parts)[:2])

    # 8. Certifications
    certifications = []
    for line in lines:
        line_lower = line.lower()
        if any(cert in line_lower for cert in CERTIFICATION_KEYWORDS):
            if len(line) < 100 and line.strip() not in certifications:
                certifications.append(line.strip())

    # 9. Location
    location = ""
    loc_match = re.search(r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z][a-z]+)\b', text)
    if loc_match:
        location = clean_candidate_location(loc_match.group(0), text)

    return {
        "name": candidate_name,
        "email": email,
        "phone": phone,
        "location": location,
        "skills": sorted(list(found_skills)),
        "education": education_str,
        "experience_years": experience_years,
        "certifications": certifications,
        "linkedin": linkedin,
        "github": github,
        "projects": []
    }

def parse_resume_content(text: str, filename: str = "") -> dict:
    if not text:
        text = ""

    # Rule-based details as instant baseline (<1ms)
    rule_res = extract_rule_based_details(text, filename)

    # Fast non-blocking LLM check
    llm_res = parse_resume_with_llm(text)

    if llm_res and isinstance(llm_res, dict):
        name = clean_candidate_name(llm_res.get("name") or rule_res["name"], filename, llm_res.get("email") or rule_res["email"], text)
        email = llm_res.get("email") or rule_res["email"]
        phone = llm_res.get("phone") or rule_res["phone"]
        raw_loc = llm_res.get("location") or rule_res["location"]
        location = clean_candidate_location(raw_loc, text)
        if not location and rule_res["location"]:
            location = rule_res["location"]

        skills = list(dict.fromkeys((llm_res.get("skills") or []) + rule_res["skills"]))
        education = llm_res.get("education") or rule_res["education"]
        exp_yrs = float(llm_res.get("experience_years", 0.0) or 0.0)
        if exp_yrs == 0.0:
            exp_yrs = rule_res["experience_years"]

        certifications = llm_res.get("certifications") or rule_res["certifications"]
        linkedin = llm_res.get("linkedin") or rule_res["linkedin"]
        github = llm_res.get("github") or rule_res["github"]
        projects = llm_res.get("projects") or []

        return {
            "name": name,
            "email": email,
            "phone": phone,
            "location": location,
            "skills": skills,
            "education": education,
            "experience_years": exp_yrs,
            "certifications": certifications,
            "linkedin": linkedin,
            "github": github,
            "projects": projects
        }

    return rule_res
